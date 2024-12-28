from fpdf import FPDF
from docx import Document
import pandas as pd
import os


def create_test_pdf(test_dir):
    """Luo PDF-testitiedosto"""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Test PDF Document", ln=1, align="C")
    pdf.cell(200, 10, txt="This is a test content for MemoryRAG", ln=1, align="L")

    pdf_path = test_dir / "test.pdf"
    pdf.output(str(pdf_path))


def create_test_docx(test_dir):
    """Luo DOCX-testitiedosto"""
    doc = Document()
    doc.add_heading("Test Document", 0)
    doc.add_paragraph("This is a test document for MemoryRAG.")
    doc.add_paragraph("It contains some test content.")

    docx_path = test_dir / "test.docx"
    doc.save(str(docx_path))


def create_test_excel(test_dir):
    """Luo Excel-testitiedosto"""
    # Luo testidataa
    data = {
        "Name": ["Test Data", "More Data", "Final Data"],
        "Value": [100, 200, 300],
        "Description": ["First", "Second", "Third"],
    }
    df = pd.DataFrame(data)

    # Tallenna Excel-tiedostoksi
    excel_path = test_dir / "test.xlsx"
    df.to_excel(str(excel_path), index=False)


def create_test_files(test_dir):
    """Luo kaikki testitiedostot"""
    create_test_pdf(test_dir)
    create_test_docx(test_dir)
    create_test_excel(test_dir)
