"""Luo testitiedostot tiedostokäsittelijöiden testejä varten"""

from pathlib import Path
import docx
from docx.shared import Pt
from reportlab.pdfgen import canvas


def create_test_files():
    """Luo PDF, DOCX ja muut testitiedostot"""
    test_dir = Path("tests/test_files")
    test_dir.mkdir(exist_ok=True)

    # Luo PDF
    pdf_path = test_dir / "test.pdf"
    c = canvas.Canvas(str(pdf_path))
    c.drawString(100, 750, "PDF Testidokumentti")
    c.drawString(100, 700, "Ensimmäinen kappale")
    c.drawString(100, 650, "Toinen kappale")
    c.save()

    # Luo DOCX
    doc = docx.Document()
    doc.add_heading("DOCX Testidokumentti", 0)
    doc.add_paragraph("Ensimmäinen kappale")
    doc.add_paragraph("Toinen kappale")
    doc.save(test_dir / "test.docx")

    # Luo HTML
    html_content = """
    <html>
        <head><title>Testisivu</title></head>
        <body>
            <h1>HTML Testidokumentti</h1>
            <p>Ensimmäinen kappale</p>
            <div>Toinen kappale</div>
            <article>Kolmas kappale</article>
            <script>console.log("Tätä ei pitäisi näkyä");</script>
            <style>.hidden { display: none; }</style>
        </body>
    </html>
    """
    (test_dir / "test.html").write_text(html_content, encoding="utf-8")

    # Luo TXT
    txt_content = """Otsikko

Ensimmäinen kappale
sisältää useita rivejä
tekstiä.

Toinen kappale on
myös monirivinenn.

Kolmas kappale."""

    (test_dir / "test.txt").write_text(txt_content, encoding="utf-8")


if __name__ == "__main__":
    create_test_files()
