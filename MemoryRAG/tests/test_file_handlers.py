"""Testit tiedostokäsittelijöille"""

import pytest
from pathlib import Path
import docx
from reportlab.pdfgen import canvas
from memoryrag.file_handlers import read_document, read_spreadsheet
from memoryrag.file_handlers.pdf_handler import read_pdf
from memoryrag.file_handlers.docx_handler import read_docx
from memoryrag.file_handlers.html_handler import read_html
import pandas as pd


@pytest.fixture(scope="session", autouse=True)
def create_test_files(tmp_path_factory):
    """Luo testitiedostot ennen testejä"""
    test_dir = tmp_path_factory.mktemp("test_files")

    # Luo PDF
    pdf_path = test_dir / "test.pdf"
    c = canvas.Canvas(str(pdf_path))
    c.drawString(100, 750, "PDF Testidokumentti")
    c.drawString(100, 700, "Ensimmäinen kappale")
    c.drawString(100, 650, "Toinen kappale")
    c.save()

    # Luo DOCX
    doc = docx.Document()
    doc.add_heading("DOCX Testidokumentti", 0)
    doc.add_paragraph("Ensimmäinen kappale")
    doc.add_paragraph("Toinen kappale")
    doc.save(test_dir / "test.docx")

    # Luo HTML
    html_content = """
    <html>
        <head><title>Testisivu</title></head>
        <body>
            <h1>HTML Testidokumentti</h1>
            <p>Ensimmäinen kappale</p>
            <div>Toinen kappale</div>
            <article>Kolmas kappale</article>
        </body>
    </html>
    """
    (test_dir / "test.html").write_text(html_content, encoding="utf-8")

    # Luo TXT
    txt_content = """Otsikko

    Ensimmäinen kappale
    sisältää useita rivejä
    tekstiä.

    Toinen kappale on
    myös monirivinenn.

    Kolmas kappale."""

    (test_dir / "test.txt").write_text(txt_content, encoding="utf-8")

    yield test_dir

    # Siivoa testitiedostot
    for f in test_dir.glob("*"):
        f.unlink()
    test_dir.rmdir()


def test_read_txt(create_test_files):
    """Testaa tekstitiedoston lukeminen"""
    test_dir = create_test_files  # Käytä fixture:n palauttamaa polkua
    paragraphs = read_document(test_dir / "test.txt")
    assert len(paragraphs) == 4
    assert "Otsikko" in paragraphs[0]
    assert "Ensimmäinen kappale" in paragraphs[1]
    assert "Toinen kappale" in paragraphs[2]
    assert "Kolmas kappale" in paragraphs[3]


def test_read_html(create_test_files):
    """Testaa HTML-tiedoston lukeminen"""
    test_dir = create_test_files
    paragraphs = read_document(test_dir / "test.html")
    assert len(paragraphs) >= 4
    assert "HTML Testidokumentti" in paragraphs[0]
    assert "Ensimmäinen kappale" in "".join(paragraphs)
    assert "Toinen kappale" in "".join(paragraphs)
    assert "Kolmas kappale" in "".join(paragraphs)


def test_read_pdf(create_test_files):
    """Testaa PDF-tiedoston lukeminen"""
    test_dir = create_test_files
    paragraphs = read_document(test_dir / "test.pdf")
    assert len(paragraphs) >= 3
    assert "PDF Testidokumentti" in "".join(paragraphs)
    assert "Ensimmäinen kappale" in "".join(paragraphs)
    assert "Toinen kappale" in "".join(paragraphs)


def test_read_docx(create_test_files):
    """Testaa Word-tiedoston lukeminen"""
    test_dir = create_test_files
    paragraphs = read_document(test_dir / "test.docx")
    assert len(paragraphs) >= 3
    assert "DOCX Testidokumentti" in paragraphs[0]
    assert "Ensimmäinen kappale" in paragraphs[1]
    assert "Toinen kappale" in paragraphs[2]


def test_read_html_url():
    """Testaa HTML:n lukeminen URL:stä"""
    url = "https://example.com/nonexistent.html"
    with pytest.raises(RuntimeError):
        read_html(url)


def test_invalid_file():
    """Testaa virheellisen tiedoston käsittely"""
    with pytest.raises(FileNotFoundError):
        read_document("olematon_tiedosto.txt")

    with pytest.raises(ValueError):
        read_document("tiedosto.invalid")


def test_empty_file(tmp_path):
    """Testaa tyhjän tiedoston käsittely"""
    empty_file = tmp_path / "empty.txt"
    empty_file.write_text("")
    paragraphs = read_document(empty_file)
    assert len(paragraphs) == 0


def test_large_file(tmp_path):
    """Testaa suuren tiedoston käsittely"""
    # Luo iso tekstitiedosto
    large_file = tmp_path / "large.txt"
    content = "Tämä on testikappale.\n\n" * 1000  # 1000 kappaletta
    large_file.write_text(content, encoding="utf-8")

    # Lue ja testaa
    paragraphs = read_document(large_file)
    assert len(paragraphs) == 1000
    assert all("Tämä on testikappale" in p for p in paragraphs)
    assert all(len(p.strip()) > 0 for p in paragraphs)  # Ei tyhjiä kappaleita


def test_special_characters(tmp_path):
    """Testaa erikoismerkkien käsittely"""
    # Luo tiedosto erikoismerkeillä
    special_file = tmp_path / "special.txt"
    content = """
    Ääkköset: äöåÄÖÅ
    
    Erikoismerkit: !@#$%^&*()
    
    Emoji: 😀🌟🎉
    
    Unicode: ∑∏∆∇∫
    """
    special_file.write_text(content, encoding="utf-8")

    # Lue ja testaa
    paragraphs = read_document(special_file)
    assert len(paragraphs) == 4
    assert "äöåÄÖÅ" in paragraphs[0]  # Ääkköset säilyvät
    assert "!@#$%^&*()" in paragraphs[1]  # Erikoismerkit säilyvät
    assert "😀🌟🎉" in paragraphs[2]  # Emojit säilyvät
    assert "∑∏∆∇∫" in paragraphs[3]  # Unicode-merkit säilyvät


def test_read_csv(create_test_files):
    """Testaa CSV-tiedoston lukeminen"""
    test_dir = create_test_files

    # Luo CSV-tiedosto
    csv_content = """nimi,ikä,kaupunki
    Matti,30,Helsinki
    Liisa,25,Tampere
    Pekka,35,Turku"""

    csv_file = test_dir / "test.csv"
    csv_file.write_text(csv_content, encoding="utf-8")

    # Lue ja testaa
    paragraphs = read_document(csv_file)
    assert len(paragraphs) == 5  # meta + otsikko + 3 riviä
    assert "nimi | ikä | kaupunki" in paragraphs[1]  # Otsikkorivi on nyt indeksissä 1
    assert "Matti | 30 | Helsinki" in paragraphs[2]  # Data alkaa indeksistä 2


def test_read_excel(create_test_files):
    """Testaa Excel-tiedoston lukeminen"""
    test_dir = create_test_files

    # Luo DataFrame
    df = pd.DataFrame(
        {
            "nimi": ["Matti", "Liisa", "Pekka"],
            "ikä": [30, 25, 35],
            "kaupunki": ["Helsinki", "Tampere", "Turku"],
        }
    )

    # Tallenna Excel-tiedostoksi
    excel_file = test_dir / "test.xlsx"
    df.to_excel(excel_file, index=False)

    # Lue ja testaa
    paragraphs = read_document(excel_file)
    assert len(paragraphs) == 5  # meta + otsikko + 3 riviä
    assert "nimi | ikä | kaupunki" in paragraphs[1]  # Otsikkorivi on nyt indeksissä 1
    assert "Matti | 30 | Helsinki" in paragraphs[2]  # Data alkaa indeksistä 2


def test_invalid_spreadsheet(tmp_path):
    """Testaa virheellisen taulukon käsittely"""
    # Tyhjä taulukko
    empty_csv = tmp_path / "empty.csv"
    empty_csv.write_text("", encoding="utf-8")
    paragraphs = read_document(empty_csv)
    assert len(paragraphs) == 0

    # Virheellinen CSV (väärä määrä sarakkeita)
    invalid_csv = tmp_path / "invalid.csv"
    invalid_csv.write_text(
        """a,b
    1,2,3
    4,5,6,7
    8,9""",
        encoding="utf-8",
    )  # Joka rivillä eri määrä sarakkeita
    with pytest.raises(RuntimeError):
        read_document(invalid_csv)

    # Liian suuri tiedosto
    large_df = pd.DataFrame(
        {
            "col": range(2000)  # Yli max_rows
        }
    )
    large_file = tmp_path / "large.xlsx"
    large_df.to_excel(large_file, index=False)
    paragraphs = read_document(large_file)
    assert len(paragraphs) <= 1002  # meta + header + 1000 rows


def test_spreadsheet_metadata(tmp_path):
    """Testaa taulukon metatietojen käsittely"""
    # Luo CSV-tiedosto metatietojen testaamiseen
    csv_content = """nimi,ikä,kaupunki,tyhjä_sarake
    Matti,30,Helsinki,
    Liisa,25,Tampere,
    Pekka,35,Turku,
    ,,,"""  # Tyhjä rivi

    csv_file = tmp_path / "metadata_test.csv"
    csv_file.write_text(csv_content, encoding="utf-8")

    # Lue ja testaa
    paragraphs = read_document(csv_file)
    meta = paragraphs[0]

    # Tarkista metatiedot
    assert "Taulukko: metadata_test.csv" in meta
    assert "Sarakkeet: 3" in meta  # Tyhjä sarake poistettu
    assert "Rivit: 3" in meta  # Tyhjä rivi poistettu

    # Tarkista sarakkeiden nimet
    assert "nimi | ikä | kaupunki" in paragraphs[1]  # Kaikki pieniä kirjaimia


def test_row_limit(tmp_path):
    """Testaa rivimäärän rajoitus"""
    # Luo suuri CSV-tiedosto
    rows = ["rivi,arvo"] + [f"{i},{i*2}" for i in range(2000)]
    csv_content = "\n".join(rows)

    csv_file = tmp_path / "large_test.csv"
    csv_file.write_text(csv_content, encoding="utf-8")

    # Testaa oletusraja (1000)
    paragraphs = read_document(csv_file)
    assert len(paragraphs) == 1002  # meta + header + 1000 riviä

    # Testaa mukautettu raja
    paragraphs = read_spreadsheet(csv_file, max_rows=500)
    assert len(paragraphs) == 502  # meta + header + 500 riviä


def test_empty_handling(tmp_path):
    """Testaa tyhjien rivien ja sarakkeiden käsittely"""
    # Luo CSV tyhjillä riveillä ja sarakkeilla
    csv_content = """nimi,ikä,kaupunki
    Matti,30,Helsinki


    Liisa,25,Tampere


    Pekka,35,Turku
    """  # Käytä vain yhtä pilkkua per sarake

    csv_file = tmp_path / "empty_test.csv"
    csv_file.write_text(csv_content, encoding="utf-8")

    # Lue ja testaa
    paragraphs = read_document(csv_file)

    # Tarkista että tyhjät rivit on poistettu
    assert len(paragraphs) == 5  # meta + header + 3 riviä


def test_column_formatting(tmp_path):
    """Testaa sarakkeiden nimien muotoilu"""
    # Luo CSV sekavilla sarakenimiä
    csv_content = """NIMI,Ikä ,  kaupunki  , Posti_Numero
    Matti,30,Helsinki,00100
    Liisa,25,Tampere,33100"""

    csv_file = tmp_path / "format_test.csv"
    csv_file.write_text(csv_content, encoding="utf-8")

    # Lue ja testaa
    paragraphs = read_document(csv_file)

    # Tarkista sarakkeiden muotoilu
    header = paragraphs[1]
    assert "nimi | ikä | kaupunki | posti_numero" in header.lower()
    assert "|  " not in header  # Ei ylimääräisiä välilyöntejä
